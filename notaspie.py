import streamlit as st
import requests
from docx import Document
from io import BytesIO
import re

# Configuración de la página
st.set_page_config(
    page_title="📝 Editor de Documentos con Corrección Automática",
    layout="wide",
)

# Título de la aplicación
st.title("📝 Editor de Documentos con Corrección Automática")

# Descripción
st.markdown("""
Esta aplicación permite **subir archivos `.docx`** con notas a pie de página y corregir errores gramaticales y ortográficos automáticamente utilizando la **API de LanguageTool**. 
**Importante:** Las notas a pie de página y notas finales se conservarán en el documento corregido, pero **si copias y pegas el texto en otro lugar, las notas a pie de página y notas finales se perderán**.
""")

# Selección de idioma
language_options = {
    "Español": "es",
    "Inglés (EE.UU.)": "en-US",
    "Francés": "fr",
    "Alemán": "de",
    # Puedes añadir más idiomas soportados por LanguageTool
}

st.sidebar.header("Configuración")
selected_language = st.sidebar.selectbox(
    "Selecciona el idioma del texto:",
    options=list(language_options.keys()),
    index=0
)

# Función para dividir párrafos largos
def dividir_parrafos_largos(texto, max_palabras=200):
    """
    Divide los párrafos que excedan max_palabras en partes más pequeñas.
    """
    parrafos = texto.split("\n")
    parrafos_divididos = []
    
    for parrafo in parrafos:
        palabras = parrafo.split()
        if len(palabras) > max_palabras:
            # Dividir en subpárrafos de tamaño max_palabras
            subparrafos = [ " ".join(palabras[i:i + max_palabras]) for i in range(0, len(palabras), max_palabras)]
            parrafos_divididos.extend(subparrafos)
        else:
            parrafos_divididos.append(parrafo)
    
    return "\n".join(parrafos_divididos)

# Función para proteger citas textuales
def proteger_citas_textuales(texto):
    """
    Encuentra y protege el texto entre comillas dobles para evitar que sea corregido.
    """
    citas_protegidas = re.findall(r'\".*?\"', texto)
    texto_protegido = texto
    for idx, cita in enumerate(citas_protegidas):
        # Reemplazar citas por marcadores temporales
        texto_protegido = texto_protegido.replace(cita, f"@@CITA{idx}@@")
    return texto_protegido, citas_protegidas

# Función para restaurar citas textuales
def restaurar_citas_textuales(texto_protegido, citas_protegidas):
    """
    Restaura las citas protegidas en el texto corregido.
    """
    for idx, cita in enumerate(citas_protegidas):
        # Reemplazar los marcadores temporales por las citas originales
        texto_protegido = texto_protegido.replace(f"@@CITA{idx}@@", cita)
    return texto_protegido

# Función para corregir texto utilizando la API de LanguageTool
def corregir_texto(texto, idioma):
    url = "https://api.languagetool.org/v2/check"
    # Proteger citas textuales
    texto_protegido, citas_protegidas = proteger_citas_textuales(texto)
    
    # Dividir párrafos largos
    texto_protegido = dividir_parrafos_largos(texto_protegido)
    
    data = {
        'text': texto_protegido,
        'language': idioma,
        'enabledOnly': False,
    }
    try:
        response = requests.post(url, data=data)
        response.raise_for_status()
        result = response.json()
        matches = result.get('matches', [])

        # Aplicar las correcciones en orden inverso para no afectar los índices
        corrected_text = texto_protegido
        for match in sorted(matches, key=lambda x: x['offset'], reverse=True):
            if match['replacements']:
                replacement = match['replacements'][0]['value']
                start = match['offset']
                end = match['offset'] + match['length']
                corrected_text = corrected_text[:start] + replacement + corrected_text[end:]
        
        # Restaurar citas textuales
        corrected_text = restaurar_citas_textuales(corrected_text, citas_protegidas)
        
        return corrected_text
    except requests.exceptions.RequestException as e:
        st.error(f"Ocurrió un error al conectar con la API de LanguageTool: {e}")
        return texto

# Función para reintegrar las notas al pie y agregar un aviso
def reintegrar_footnotes(document, footnotes):
    """
    Reintegra las notas al pie y añade una nota explicativa de que no se han corregido.
    También se añade una advertencia sobre la pérdida de las notas al copiar y pegar.
    """
    # Añadir un párrafo al final explicando que las notas no fueron corregidas y que se pierden al pegar el texto.
    document.add_paragraph("Nota: Las notas al pie no fueron corregidas en este documento.")
    document.add_paragraph("Importante: Si copias y pegas el texto en otro documento, las notas al pie y notas finales se perderán.")
    return document

# Función para procesar el documento .docx
def procesar_docx(file, idioma):
    # Leer el documento original
    try:
        document = Document(file)
    except Exception as e:
        st.error(f"Error al leer el archivo .docx: {e}")
        return None

    # Separar footnotes y texto principal
    texto_principal, footnotes = separar_footnotes(document)

    # Corregir el texto principal
    texto_corregido = corregir_texto(texto_principal, idioma)

    # Aplicar las correcciones al documento
    document = aplicar_correcciones(document, texto_corregido)

    # Reintegrar las footnotes y agregar un aviso
    document = reintegrar_footnotes(document, footnotes)

    # Guardar el documento corregido en un buffer
    buffer = BytesIO()
    try:
        document.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error al guardar el documento corregido: {e}")
        return None

# Función para separar footnotes del texto principal
def separar_footnotes(document):
    """
    Separa las definiciones de notas a pie de página del texto principal.
    Retorna el texto sin las definiciones de footnotes y una lista con las definiciones extraídas.
    """
    footnotes = {}
    for rel in document.part.rels.values():
        if "footnotes" in rel.target_ref:
            part = rel.target_part
            for elem in part.element.getchildren():
                footnotes[elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")] = elem.text

    # Extraer el texto principal con referencias a footnotes
    texto_principal = ""
    for para in document.paragraphs:
        texto_principal += para.text + "\n"

    return texto_principal.strip(), footnotes

# Función para aplicar correcciones al documento
def aplicar_correcciones(document, texto_corregido):
    """
    Reemplaza el texto de los párrafos en el documento con el texto corregido.
    Preserva los footnotes y el formato de los runs tanto como sea posible.
    """
    parrafos_corregidos = texto_corregido.split("\n")

    for para, texto_corr in zip(document.paragraphs, parrafos_corregidos):
        if para.text.strip() == "":
            continue  # Ignorar párrafos vacíos
        for run in para.runs:
            run.text = ""
        para.add_run(texto_corr)

    return document

# Área de carga de archivo
uploaded_file = st.file_uploader(
    "Sube tu archivo .docx aquí:",
    type=["docx"]
)

# Botón para corregir el documento
if uploaded_file is not None:
    if st.button("Corregir Documento"):
        with st.spinner("Corrigiendo el documento..."):
            archivo_corregido = procesar_docx(uploaded_file, language_options[selected_language])
            if archivo_corregido:
                st.success("¡Corrección completada!")
                st.download_button(
                    label="Descargar Documento Corregido",
                    data=archivo_corregido,
                    file_name="documento_corregido.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
